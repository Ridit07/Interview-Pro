import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from moviepy.editor import VideoFileClip
from docx import Document
import cv2
from facenet_pytorch import MTCNN, InceptionResnetV1
import numpy as np
import ssl
ssl._create_default_https_context = ssl._create_unverified_context
import requests
import random
import os
from PIL import Image
from Fashion_detection import process_images
from docx.shared import Inches
import tensorflow as tf
import shutil
from emotion import analyze_emotions
from voice import main_audio
from gaze_final import analyze_gaze_in_video
from flask import Flask, request, jsonify
import os
import pickle
from toxic import load_model, preprocess  # Import from your toxicity detection code
import pandas as pd
from tensorflow.keras.preprocessing.sequence import pad_sequences

app = Flask(__name__)

# Initialize the Whisper model and other settings
model_id = "openai/whisper-large-v3"
device = "cuda:0" if torch.cuda.is_available() else "cpu"
torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

model = AutoModelForSpeechSeq2Seq.from_pretrained(
    model_id,
    torch_dtype=torch_dtype,
    low_cpu_mem_usage=True,
    use_safetensors=True
)
model.to(device)
processor = AutoProcessor.from_pretrained(model_id)

pipe = pipeline(
    "automatic-speech-recognition",
    model=model,
    tokenizer=processor.tokenizer,
    feature_extractor=processor.feature_extractor,
    max_new_tokens=128,
    chunk_length_s=15,
    batch_size=16,
    return_timestamps=True,
    torch_dtype=torch_dtype,
    device=device,
)

TOKENIZER_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/tokenizer.pkl'
MODEL_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/model.pth'
EMBEDDING_MATRIX_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/embedding_matrix.npy'

embedding_matrix = np.load(EMBEDDING_MATRIX_PATH)
with open(TOKENIZER_PATH, 'rb') as f:
    tokenizer = pickle.load(f)

toxicity_model = load_model(MODEL_PATH, embedding_matrix)

@app.route('/generate_report', methods=['POST'])
def generate_report():
    # Get the interviewer_id from the request
    # data = request.get_json()
    # interviewer_id = data.get("interviewer_id")

    # if not interviewer_id:
    #     return jsonify({"error": "interviewer_id is required"}), 400

    # # Define paths based on interviewer_id
    # base_path = "/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/database/static/uploads/saved_files"
    # video_path = os.path.join(base_path, f"interview_rec/{interviewer_id}/final.mp4")
    # audio_path = os.path.join(base_path, f"interview_rec/{interviewer_id}/final.mp3")
    # report_path = os.path.join(base_path, f"interviewer_report/ailogfile_{interviewer_id}.docx")

    video_path = f"/Users/riditjain/Downloads/akjv.mp4"
    audio_path = f"/Users/riditjain/Downloads/akjv.mp3"
    report_path = f"/Users/riditjain/Downloads/akjv.docx"

    # Ensure the video file exists
    if not os.path.exists(video_path):
        return jsonify({"error": "Video file not found"}), 404

    # Video processing
    video_clip = VideoFileClip(video_path)
    audio_clip = video_clip.audio
    audio_clip.write_audiofile(audio_path, codec='mp3')
    video_duration = video_clip.duration  # Total duration of the video in seconds

    # Transcribe audio
    result = pipe(audio_path)
    transcribed_text = result['text']

    # Send transcription to NLP server
    nlp_response = requests.post("http://127.0.0.1:5002/analyze", json={"text": transcribed_text})

    # Initialize Word document with better formatting
    doc = Document()
    doc.add_heading('Interview Analysis Report', 0)

    # Transcribed Responses Section
    doc.add_heading('Transcribed Responses', level=1)
    doc.add_paragraph(transcribed_text)

    # Face Detection and Tracking Setup
    device = torch.device('cuda:0' if torch.cuda.is_available() else 'cpu')
    mtcnn = MTCNN(keep_all=True, device=device)
    resnet = InceptionResnetV1(pretrained='vggface2').eval().to(device)
    video_capture = cv2.VideoCapture(video_path)

    persistent_embeddings = {}
    next_id = 1
    cheating_log = []
    frame_rate = video_capture.get(cv2.CAP_PROP_FPS)
    person1_last_seen = None
    person1_out_of_frame_start = None

    REENTRY_FRAME_THRESHOLD = 30  # ~1 second at 30 FPS
    last_seen_frame = None

    def find_matching_id(new_embedding, embeddings, threshold=0.8):
        if len(embeddings) == 0:
            return None
        min_distance = float('inf')
        matching_id = None
        for face_id, emb in embeddings.items():
            distance = torch.nn.functional.pairwise_distance(new_embedding, emb).min().item()
            if distance < min_distance:
                min_distance = distance
                matching_id = face_id
        if min_distance < threshold:
            return matching_id
        return None

    frame_number = 0
    while True:
        ret, frame = video_capture.read()
        if not ret:
            break

        timestamp = frame_number / frame_rate
        rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        boxes, _ = mtcnn.detect(rgb_frame)

        if boxes is not None:
            for box in boxes:
                x1, y1, x2, y2 = [max(0, int(coord)) for coord in box]
                x2 = min(x2, rgb_frame.shape[1])
                y2 = min(y2, rgb_frame.shape[0])

                face = rgb_frame[y1:y2, x1:x2]
                if face.size == 0:
                    continue

                face = cv2.resize(face, (160, 160))
                face = torch.tensor(face).permute(2, 0, 1).float().div(255).unsqueeze(0).to(device)

                with torch.no_grad():
                    embedding = resnet(face).detach()

                matching_id = find_matching_id(embedding, persistent_embeddings)

                if matching_id is None:
                    if timestamp < video_duration - 1:  # Ignore last second entries
                        if next_id == 1:
                            person1_last_seen = timestamp  # Set first person as Candidate
                        persistent_embeddings[next_id] = embedding
                        cheating_log.append(f"Person {next_id} enters at {timestamp:.2f} seconds.")
                        next_id += 1
                else:
                    if matching_id == 1:
                        person1_last_seen = timestamp
                        last_seen_frame = frame_number
                        person1_out_of_frame_start = None

        # Check if Candidate has been missing for more than 2 seconds
        if person1_last_seen is not None and last_seen_frame is not None:
            frames_since_seen = frame_number - last_seen_frame
            if frames_since_seen > (2 * frame_rate):  # 2 seconds in frames
                if person1_out_of_frame_start is None:
                    person1_out_of_frame_start = person1_last_seen
                    cheating_log.append(f"Candidate left the camera at {person1_out_of_frame_start:.2f} seconds for more than 2 seconds.")

        frame_number += 1

    video_capture.release()
    cv2.destroyAllWindows()

    # Cheating Detection Section
    doc.add_heading('Cheating Detection Summary', level=1)
    if cheating_log:
        for log_entry in cheating_log:
            doc.add_paragraph(log_entry)
    else:
        doc.add_paragraph("No signs of cheating detected during the interview.")

    # NLP Insights
    if nlp_response.status_code == 200:
        analysis_data = nlp_response.json()

        doc.add_heading('Linguistic and Content Analysis', level=1)

        # Thought Process Metrics
        thought_process = analysis_data['thought_process_metrics']
        doc.add_heading('Thought Process Metrics', level=2)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'LightShading-Accent1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Total Sentences'
        row_cells[1].text = str(thought_process['total_sentences'])
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Average Sentence Length'
        row_cells[1].text = f"{thought_process['average_sentence_length']:.2f} words"
        row_cells = table.rows[3].cells
        row_cells[0].text = 'Logical Connector Count'
        row_cells[1].text = str(thought_process['logical_connector_count'])

        # Keywords
        doc.add_heading('Identified Keywords', level=2)
        doc.add_paragraph(", ".join(analysis_data['keywords']))

        # Topics
        doc.add_heading('Topics Discussed', level=2)
        doc.add_paragraph(
            "The following topics were identified in the candidate's responses, indicating areas of focus during the interview."
        )
        for topic_num, topic_content in enumerate(analysis_data['topics']):
            # Extract keywords from the topic string and remove probabilities
            topic_keywords = ", ".join(
                [word.split("*")[1].replace('"', '').strip() for word in topic_content[1].split(" + ")]
            )
            doc.add_paragraph(f"Topic {topic_num + 1}: {topic_keywords}")

        # Named Entities
        doc.add_heading('Mentioned Entities', level=2)
        for entity in analysis_data['entities']:
            doc.add_paragraph(f"{entity[0]} ({entity[1]})")

        # Complexity Score
        doc.add_heading('Readability Score', level=2)
        complexity_score = analysis_data['complexity_score']
        doc.add_paragraph(
            f"The candidate's responses have a Flesch-Kincaid grade level of {complexity_score:.2f}, "
            "indicating the readability and complexity of the language used."
        )

        # Technology Experience
        doc.add_heading('Technologies and Skills Mentioned', level=2)
        if analysis_data['technology_experience']:
            for experience in analysis_data['technology_experience']:
                doc.add_paragraph(f"- {experience}")
        else:
            doc.add_paragraph("No specific technologies or skills were mentioned.")

    # Attire Detection Section
    def extract_random_frames(video_path, output_dir, num_frames=5):
        video_clip = VideoFileClip(video_path)
        duration = video_clip.duration
        timestamps = sorted(random.sample([random.uniform(0, duration) for _ in range(num_frames)], num_frames))
        image_paths = []
        for idx, t in enumerate(timestamps):
            frame = video_clip.get_frame(t)
            frame_image = Image.fromarray(frame)
            image_path = os.path.join(output_dir, f"frame_{idx+1}.png")
            frame_image.save(image_path)
            image_paths.append(image_path)
        return image_paths

    # Define directories
    extracted_frames_dir = "/Users/riditjain/Downloads/extracted_frames"
    processed_images_dir = "/Users/riditjain/Downloads/processed_attire_images"
    os.makedirs(extracted_frames_dir, exist_ok=True)
    os.makedirs(processed_images_dir, exist_ok=True)

    # Extract and process frames
    random_image_paths = extract_random_frames(video_path, extracted_frames_dir, num_frames=5)
    processed_image_paths = process_images(random_image_paths, processed_images_dir, threshold=0.8)

    doc.add_heading('Professional Attire Assessment', level=1)
    if processed_image_paths:
        doc.add_paragraph("The following images show the candidate's attire during the interview:")
        for img_path in processed_image_paths:
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
            else:
                doc.add_paragraph(f"Image {img_path} not found.")
    else:
        doc.add_paragraph("No images available for attire assessment.")

    # Emotion Detection Section
    doc.add_heading('Emotional Tone Analysis', level=1)

    # Analyze emotions in the video
    emotion_percentages = analyze_emotions(video_path)

    # Map emotions to more professional terms
    emotion_mapping = {
        'happy': 'Positive',
        'sad': 'Negative',
        'angry': 'Frustrated',
        'surprise': 'Surprised',
        'Neutral': 'Neutral',
        'fearful': 'Anxious',
        'disgust': 'Displeased'
    }

    # Write emotion detection results to the Word document
    if emotion_percentages:
        doc.add_paragraph("The emotional expressions detected during the interview are as follows:")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'LightShading-Accent2'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Emotion'
        hdr_cells[1].text = 'Percentage'
        for emotion, percentage in emotion_percentages.items():
            mapped_emotion = emotion_mapping.get(emotion, emotion)
            row_cells = table.add_row().cells
            row_cells[0].text = mapped_emotion
            row_cells[1].text = f"{percentage:.2f}%"
    else:
        doc.add_paragraph("No significant emotional expressions detected.")

    # Toxicity Detection Section
    def get_toxicity_predictions(text, model, tokenizer, max_len=220):
        text = pd.Series([text])
        text = preprocess(text)
        sequences = tokenizer.texts_to_sequences(text)
        padded_sequence = pad_sequences(sequences, maxlen=max_len)
        input_tensor = torch.tensor(padded_sequence, dtype=torch.long)

        with torch.no_grad():
            output = model(input_tensor)
            probabilities = torch.sigmoid(output).cpu().numpy()
        return probabilities[0]

    # Obtain toxicity prediction results
    toxicity_scores = get_toxicity_predictions(transcribed_text, toxicity_model, tokenizer)

    # Convert toxicity prediction to a readable format
    def format_toxicity_results(prediction):
        labels = [
            "Appropriate Language",
            "Highly Inappropriate Language",
            "Use of Obscene Language",
            "Identity Attacks",
            "Insults",
            "Threats"
        ]
        result_text = ""
        for label, score in zip(labels, prediction):
            result_text += f"{label}: {score * 100:.2f}%\n"
        return result_text

    toxicity_results_text = format_toxicity_results(toxicity_scores)

    doc.add_heading('Language Appropriateness Analysis', level=1)
    doc.add_paragraph("The following percentages indicate the likelihood of inappropriate language usage:")
    doc.add_paragraph(toxicity_results_text)

    # Speaker Diarization Section
    speaker_result = main_audio(audio_path)
    doc.add_heading('Speaker Analysis', level=1)
    doc.add_paragraph(speaker_result)

    # Gaze Detection Section
    left_count, right_count, long_left_count, long_right_count = analyze_gaze_in_video(video_path)
    doc.add_heading('Eye Gaze Analysis', level=1)
    gaze_table = doc.add_table(rows=1, cols=2)
    gaze_table.style = 'LightShading-Accent3'
    gaze_hdr_cells = gaze_table.rows[0].cells
    gaze_hdr_cells[0].text = 'Metric'
    gaze_hdr_cells[1].text = 'Count'
    gaze_data = [
        ('Total times looking left', left_count),
        ('Total times looking right', right_count),
        ('Times looking left for more than 5 seconds', long_left_count),
        ('Times looking right for more than 5 seconds', long_right_count)
    ]
    for metric, count in gaze_data:
        row_cells = gaze_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(count)

    # Lip Sync Detection Section
    flask_server_url = "http://127.0.0.1:5003/lip_sync_check"

    def get_lip_sync_percentage(video_path):
        try:
            response = requests.post(
                flask_server_url,
                json={"video_path": video_path}
            )
            response_data = response.json()
            if response.status_code == 200:
                return response_data.get("similarity_percentage", "N/A")
            else:
                print("Error from lip-sync detection server:", response_data.get("error", "Unknown error"))
                return "N/A"
        except Exception as e:
            print("Failed to connect to lip-sync detection server:", e)
            return "N/A"

    lip_sync_percentage = get_lip_sync_percentage(video_path)

    doc.add_heading('Lip Sync Consistency', level=1)
    doc.add_paragraph(
        f"The lip-sync similarity between the audio and video is {lip_sync_percentage}. "
        "This measures the synchronization between spoken words and lip movements."
    )

    # Save the Word document
    doc.save(report_path)

    # Clean up temporary directories
    shutil.rmtree(extracted_frames_dir, ignore_errors=True)
    shutil.rmtree(processed_images_dir, ignore_errors=True)

    # try:
    #     response = requests.post(
    #         "http://127.0.0.1:5000/api/update_status_and_ai_report",
    #         json={"interview_id": interviewer_id}
    #     )
    #     if response.status_code == 200:
    #         print("Status and AI Report path updated successfully.")
    #     else:
    #         print(f"Failed to update status and AI report path: {response.json().get('error')}")
    # except Exception as e:
    #     print(f"Error during status and AI report path update: {e}")

    # return jsonify({"success": True, "report_path": report_path}), 200

# Uncomment the following line if you want to run generate_report directly
generate_report()

if __name__ == "__main__":
    app.run(debug=True, port=5007)
