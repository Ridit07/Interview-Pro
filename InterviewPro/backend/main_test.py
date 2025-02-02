import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from moviepy.editor import VideoFileClip
from docx import Document
import cv2
from facenet_pytorch import MTCNN, InceptionResnetV1
import numpy as np
import ssl
ssl._create_default_https_context = ssl._create_unverified_context
import requests
import random
import os
from PIL import Image
from Fashion_detection import process_images
from docx.shared import Inches
import tensorflow as tf
import shutil
from emotion import analyze_emotions
from voice import main_audio
from gaze_final import analyze_gaze_in_video

# Initialize the Whisper model and other settings
model_id = "openai/whisper-large-v3"
device = "cuda:0" if torch.cuda.is_available() else "cpu"
torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

model = AutoModelForSpeechSeq2Seq.from_pretrained(model_id, torch_dtype=torch_dtype, low_cpu_mem_usage=True, use_safetensors=True)
model.to(device)
processor = AutoProcessor.from_pretrained(model_id)

pipe = pipeline(
    "automatic-speech-recognition",
    model=model,
    tokenizer=processor.tokenizer,
    feature_extractor=processor.feature_extractor,
    max_new_tokens=128,
    chunk_length_s=15,
    batch_size=16,
    return_timestamps=True,
    torch_dtype=torch_dtype,
    device=device,
)

# Video processing
video_path = "/Users/riditjain/Downloads/lavanya_video.mp4"
audio_path = "/Users/riditjain/Downloads/lavanya_video.mp3"
video_clip = VideoFileClip(video_path)
audio_clip = video_clip.audio
audio_clip.write_audiofile(audio_path, codec='mp3')
video_duration = video_clip.duration  # Total duration of the video in seconds

# Transcribe audio
result = pipe(audio_path)
transcribed_text = result['text']

# Send transcription to NLP server
nlp_response = requests.post("http://127.0.0.1:5002/analyze", json={"text": transcribed_text})

# Initialize Word document
doc = Document()
doc.add_heading('Translated Text', level=1)
doc.add_paragraph(transcribed_text)

# Face Detection and Tracking Setup
device = torch.device('cuda:0' if torch.cuda.is_available() else 'cpu')
mtcnn = MTCNN(keep_all=True, device=device)
resnet = InceptionResnetV1(pretrained='vggface2').eval().to(device)
video_capture = cv2.VideoCapture(video_path)

persistent_embeddings = {}
next_id = 1
cheating_log = []
frame_rate = video_capture.get(cv2.CAP_PROP_FPS)
person1_last_seen = None
person1_out_of_frame_start = None

REENTRY_FRAME_THRESHOLD = 30  # ~1 second at 30 FPS
last_seen_frame = None

def find_matching_id(new_embedding, embeddings, threshold=0.8):
    if len(embeddings) == 0:
        return None
    min_distance = float('inf')
    matching_id = None
    for face_id, emb in embeddings.items():
        distance = torch.nn.functional.pairwise_distance(new_embedding, emb).min().item()
        if distance < min_distance:
            min_distance = distance
            matching_id = face_id
    if min_distance < threshold:
        return matching_id
    return None

frame_number = 0
while True:
    ret, frame = video_capture.read()
    if not ret:
        break

    timestamp = frame_number / frame_rate
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    boxes, _ = mtcnn.detect(rgb_frame)

    if boxes is not None:
        for box in boxes:
            x1, y1, x2, y2 = [max(0, int(coord)) for coord in box]
            x2 = min(x2, rgb_frame.shape[1])
            y2 = min(y2, rgb_frame.shape[0])
            
            face = rgb_frame[y1:y2, x1:x2]
            if face.size == 0:
                continue
            
            face = cv2.resize(face, (160, 160))
            face = torch.tensor(face).permute(2,0,1).float().div(255).unsqueeze(0).to(device)
            
            with torch.no_grad():
                embedding = resnet(face).detach()
            
            matching_id = find_matching_id(embedding, persistent_embeddings)

            if matching_id is None:
                if timestamp < video_duration - 1:  # Ignore last second entries
                    if next_id == 1:
                        person1_last_seen = timestamp  # Set first person as Person 1
                    persistent_embeddings[next_id] = embedding
                    cheating_log.append(f"Person {next_id} enters at {timestamp:.2f} seconds.")
                    next_id += 1
            else:
                if matching_id == 1:
                    person1_last_seen = timestamp
                    last_seen_frame = frame_number
                    person1_out_of_frame_start = None

    # Check if Person 1 has been missing for more than 2 seconds
    if person1_last_seen is not None and last_seen_frame is not None:
        frames_since_seen = frame_number - last_seen_frame
        if frames_since_seen > (2 * frame_rate):  # 2 seconds in frames
            if person1_out_of_frame_start is None:
                person1_out_of_frame_start = person1_last_seen
                cheating_log.append(f"Person 1 left the camera at {person1_out_of_frame_start:.2f} seconds for more than 2 seconds.")

    frame_number += 1

video_capture.release()
cv2.destroyAllWindows()

# Write the Cheating Log or indicate no cheating detected
doc.add_heading('Cheating Detection Log', level=1)
if cheating_log:
    for log_entry in cheating_log:
        doc.add_paragraph(log_entry)
else:
    doc.add_paragraph("No cheating detected.")

# Add NLP insights if the response is successful
if nlp_response.status_code == 200:
    analysis_data = nlp_response.json()
    doc.add_heading('NLP Insight Analysis', level=1)
    
    # Add Thought Process Metrics
    thought_process = analysis_data['thought_process_metrics']
    doc.add_heading('Thought Process Metrics', level=2)
    doc.add_paragraph(f"Total Sentences: {thought_process['total_sentences']}")
    doc.add_paragraph(f"Average Sentence Length: {thought_process['average_sentence_length']:.2f}")
    doc.add_paragraph(f"Logical Connector Count: {thought_process['logical_connector_count']}")
    
    # Add Keywords
    doc.add_heading('Keywords', level=2)
    doc.add_paragraph(", ".join(analysis_data['keywords']))
    
    # Add Topics
    doc.add_heading('Topics', level=2)
    for topic in analysis_data['topics']:
        doc.add_paragraph(str(topic))
    
    # Add Named Entities
    doc.add_heading('Named Entities', level=2)
    for entity in analysis_data['entities']:
        doc.add_paragraph(f"{entity[0]}: {entity[1]}")
    
    # Add Complexity Score
    doc.add_heading('Complexity Score', level=2)
    doc.add_paragraph(str(analysis_data['complexity_score']))
    
    # Add Technology Experience
    doc.add_heading('Technology Experience', level=2)
    for experience in analysis_data['technology_experience']:
        doc.add_paragraph(experience)

# 3. Extract 5 Random Frames from the Video
def extract_random_frames(video_path, output_dir, num_frames=5):
    video_clip = VideoFileClip(video_path)
    duration = video_clip.duration
    timestamps = sorted(random.sample([random.uniform(0, duration) for _ in range(num_frames)], num_frames))
    image_paths = []
    for idx, t in enumerate(timestamps):
        frame = video_clip.get_frame(t)
        frame_image = Image.fromarray(frame)
        image_path = os.path.join(output_dir, f"frame_{idx+1}.png")
        frame_image.save(image_path)
        image_paths.append(image_path)
    return image_paths

# Define directories
extracted_frames_dir = "/Users/riditjain/Downloads/extracted_frames"
processed_images_dir = "/Users/riditjain/Downloads/processed_attire_images"
os.makedirs(extracted_frames_dir, exist_ok=True)
os.makedirs(processed_images_dir, exist_ok=True)

# Extract 5 random frames
random_image_paths = extract_random_frames(video_path, extracted_frames_dir, num_frames=5)

# Process the extracted frames
processed_image_paths = process_images(random_image_paths, processed_images_dir, threshold=0.8)

# 4. Insert Processed Images into the Word Document
doc.add_heading('Attire Detection', level=1)
for img_path in processed_image_paths:
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))  # Adjust width as needed
        doc.add_paragraph(os.path.basename(img_path))  # Optionally add image name as caption
    else:
        doc.add_paragraph(f"Image {img_path} not found.")

## Emotion Detection

doc.add_heading('Emotion Detection', level=1)

# Call the emotionsss function to analyze emotions in the video
emotion_percentages = analyze_emotions(video_path)

# Write emotion detection results to the Word document
if emotion_percentages:
    for emotion, percentage in emotion_percentages.items():
        doc.add_paragraph(f"{emotion}: {percentage:.2f}%")
else:
    doc.add_paragraph("No emotions detected.")

## Toxicity Detection

import pickle
from toxic import load_model,preprocess # Import from your toxicity detection code
import pandas as pd
from tensorflow.keras.preprocessing.sequence import pad_sequences

# Load the Toxicity model and tokenizer
TOKENIZER_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/tokenizer.pkl'
MODEL_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/model.pth'
EMBEDDING_MATRIX_PATH = '/Users/riditjain/Downloads/Interview-Pro-main/InterviewPro/backend/models/toxic/embedding_matrix.npy'

embedding_matrix = np.load(EMBEDDING_MATRIX_PATH)
with open(TOKENIZER_PATH, 'rb') as f:
    tokenizer = pickle.load(f)

toxicity_model = load_model(MODEL_PATH, embedding_matrix)

# Define function for properly processing the text and getting toxicity predictions
def get_toxicity_predictions(text, model, tokenizer, max_len=220):
    # Preprocess text for prediction
    text = pd.Series([text])
    text = preprocess(text)
    sequences = tokenizer.texts_to_sequences(text)
    padded_sequence = pad_sequences(sequences, maxlen=max_len)
    input_tensor = torch.tensor(padded_sequence, dtype=torch.long)
    
    # Predict and apply sigmoid for final probability scores
    with torch.no_grad():
        output = model(input_tensor)
        probabilities = torch.sigmoid(output).cpu().numpy()  # Sigmoid activation for probabilities
    return probabilities[0]  # Return the array of probabilities for the given classes

# Obtain toxicity prediction results
toxicity_scores = get_toxicity_predictions(transcribed_text, toxicity_model, tokenizer)

# Convert toxicity prediction to a readable format
def format_toxicity_results(prediction):
    labels = ["Toxicity", "Severe Toxicity", "Obscene", "Identity Attack", "Insult", "Threat"]
    result_text = ""
    for label, score in zip(labels, prediction):
        result_text += f"{label}: {score:.2f}\n"
    return result_text

toxicity_results_text = format_toxicity_results(toxicity_scores)

# Add the toxicity results to the Word document
doc.add_heading('Toxicity Detection Results', level=1)
doc.add_paragraph(toxicity_results_text)

## voice detection

speaker_result = main_audio(audio_path)

# Add the speaker detection result to the Word document
doc.add_heading('Speaker Diarization', level=1)
doc.add_paragraph(speaker_result)



## Gaze Detection

left_count, right_count, long_left_count, long_right_count =  analyze_gaze_in_video(video_path)
doc.add_heading('Gaze Detection', level=1)
doc.add_paragraph(f"Total times looking left: {left_count}")
doc.add_paragraph(f"Total times looking right: {right_count}")
doc.add_paragraph(f"Times looking left for more than 5 seconds: {long_left_count}")
doc.add_paragraph(f"Times looking right for more than 5 seconds: {long_right_count}")


flask_server_url = "http://127.0.0.1:5003/lip_sync_check"

# Send a POST request to the Flask server for lip-sync detection
def get_lip_sync_percentage(video_path):
    try:
        response = requests.post(
            flask_server_url,
            json={"video_path": video_path}
        )
        response_data = response.json()
        if response.status_code == 200:
            return response_data.get("similarity_percentage", "N/A")
        else:
            print("Error from lip-sync detection server:", response_data.get("error", "Unknown error"))
            return "N/A"
    except Exception as e:
        print("Failed to connect to lip-sync detection server:", e)
        return "N/A"

# Call the lip-sync detection server and get the similarity percentage
lip_sync_percentage = get_lip_sync_percentage(video_path)

# Add Lip Sync Detection result to the Word document
doc.add_heading('Lip Sync Detection', level=1)
doc.add_paragraph(f"Lip Sync Similarity: {lip_sync_percentage}")

# Save the Word document
doc.save('/Users/riditjain/Downloads/logfilelavi.docx')

# Clean up temporary directories if used
shutil.rmtree(extracted_frames_dir, ignore_errors=True)
shutil.rmtree(processed_images_dir, ignore_errors=True)